from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
from groq import Groq
import json, re, os, io, uuid, time
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'mcq-solver-2024')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
GEMINI_KEYS = [k.strip() for k in os.environ.get('GEMINI_API_KEYS', '').split(',') if k.strip()]
GROQ_KEYS = [k.strip() for k in os.environ.get('GROQ_API_KEYS', '').split(',') if k.strip()]
gemini_index = 0
groq_index = 0
ai_count = 0
if GEMINI_KEYS:
    ai_count += 1
if GROQ_KEYS:
    ai_count += 1
print(f"AI:{ai_count} Keys:{len(GEMINI_KEYS)+len(GROQ_KEYS)}")
sessions = {}

def xjson(text, as_list=False):
    if not text:
        return [] if as_list else {}
    try:
        text = text.strip()
        m = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
        if m:
            text = m.group(1)
        if as_list:
            m2 = re.search(r'\[[\s\S]*\]', text)
            if m2:
                return json.loads(m2.group())
        else:
            m2 = re.search(r'\{[\s\S]*\}', text)
            if m2:
                return json.loads(m2.group())
    except Exception:
        pass
    if as_list:
        return []
    return {"explanation": text[:1000] if text else "", "sources": ["তথ্যসূত্র যাচাই প্রয়োজন"], "confidence": 0.4}

def cses():
    now = time.time()
    for s in [s for s, d in sessions.items() if now - d.get('t', 0) > 3600]:
        del sessions[s]

def agemini(prompt):
    global gemini_index
    if not GEMINI_KEYS:
        return None
    for _ in range(len(GEMINI_KEYS)):
        key = GEMINI_KEYS[gemini_index % len(GEMINI_KEYS)]
        gemini_index += 1
        try:
            genai.configure(api_key=key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            r = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=0.2, max_output_tokens=2500))
            return r.text
        except Exception:
            time.sleep(0.3)
    return None

def agroq(prompt):
    global groq_index
    if not GROQ_KEYS:
        return None
    for _ in range(len(GROQ_KEYS)):
        key = GROQ_KEYS[groq_index % len(GROQ_KEYS)]
        groq_index += 1
        try:
            client = Groq(api_key=key)
            r = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": "তুমি বাংলাদেশের শিক্ষক। বাংলায় JSON এ উত্তর দাও।"}, {"role": "user", "content": prompt}], temperature=0.2, max_tokens=2500)
            return r.choices[0].message.content
        except Exception:
            time.sleep(0.3)
    return None

def pmcqs(raw):
    p = "নিচের টেক্সট থেকে MCQ extract করো। ৪ অপশন+১ উত্তর আছে এমন প্রশ্ন নাও। ওয়াটারমার্ক/প্রমো বাদ। JSON array: [{\"question_no\":1,\"question\":\"?\",\"option_a\":\"a\",\"option_b\":\"b\",\"option_c\":\"c\",\"option_d\":\"d\",\"correct_answer\":\"ক\",\"correct_option_text\":\"t\"}]\nটেক্সট:\n" + raw
    for fn in [agemini, agroq]:
        r = fn(p)
        if r:
            x = xjson(r, as_list=True)
            if x:
                return x
        time.sleep(0.3)
    return []

def pimg(path):
    if not GEMINI_KEYS:
        return []
    try:
        import PIL.Image
        img = PIL.Image.open(path)
        genai.configure(api_key=GEMINI_KEYS[0])
        v = genai.GenerativeModel('gemini-2.0-flash')
        r = v.generate_content(["ছবি থেকে MCQ extract করো। JSON array দাও।", img])
        return xjson(r.text, as_list=True)
    except Exception:
        return []

def gexp(mcq, subj=""):
    s = ("\nবিষয়: " + subj) if subj else ""
    p = "তুমি বাংলাদেশের শিক্ষক।" + s + "\nপ্রশ্ন: " + mcq.get('question', '') + "\n(ক) " + mcq.get('option_a', '') + "\n(খ) " + mcq.get('option_b', '') + "\n(গ) " + mcq.get('option_c', '') + "\n(ঘ) " + mcq.get('option_d', '') + "\nসঠিক: (" + mcq.get('correct_answer', '') + ") " + mcq.get('correct_option_text', '') + "\n\n১.সঠিক কেন ২.বাকি কেন ভুল ৩.বাংলাদেশ তথ্য ৪.২টি সোর্স(BARI,BRRI,BINA,AIS,BAU,বাংলাপিডিয়া,BBS,NCTB)\nJSON:{\"explanation\":\"...\",\"sources\":[\"s1\",\"s2\"],\"confidence\":0.9,\"key_point\":\"p\"}"
    res = {}
    g = agemini(p)
    if g:
        res['g'] = xjson(g)
    time.sleep(0.5)
    q = agroq(p)
    if q:
        res['q'] = xjson(q)
    if not res:
        return {"explanation": "ব্যাখ্যা তৈরি যায়নি।", "sources": ["API Key চেক করুন"], "key_point": ""}
    asrc = []
    best = None
    bsc = -1
    for n, r in res.items():
        if not isinstance(r, dict):
            continue
        for ss in r.get('sources', []):
            if ss and ss not in asrc:
                asrc.append(ss)
        sc = r.get('confidence', 0.5)
        if len(r.get('explanation', '')) > 200:
            sc += 0.1
        if n == 'g':
            sc += 0.05
        if sc > bsc:
            bsc = sc
            best = r
    if best:
        best['sources'] = asrc[:5]
        if len(best['sources']) < 2:
            best['sources'].extend(["বাংলাপিডিয়া", "জাতীয় তথ্য বাতায়ন"])
            best['sources'] = best['sources'][:4]
        return best
    return {"explanation": "ব্যাখ্যা পাওয়া যায়নি", "sources": asrc[:3], "key_point": ""}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/status')
def api_status():
    return jsonify({"gemini": len(GEMINI_KEYS) > 0, "groq": len(GROQ_KEYS) > 0, "together": False, "total_ai": ai_count, "total_keys": len(GEMINI_KEYS) + len(GROQ_KEYS)})

@app.route('/api/parse-text', methods=['POST'])
def api_parse_text():
    try:
        cses()
        d = request.get_json()
        t = d.get('text', '')
        if not t.strip():
            return jsonify({"error": "টেক্সট দিন"}), 400
        m = pmcqs(t)
        if not m:
            return jsonify({"error": "MCQ পাওয়া যায়নি"}), 400
        sid = str(uuid.uuid4())[:8]
        sessions[sid] = {'mcqs': m, 'header': {}, 't': time.time()}
        return jsonify({"session_id": sid, "mcqs": m, "count": len(m)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/parse-file', methods=['POST'])
def api_parse_file():
    try:
        cses()
        if 'file' not in request.files:
            return jsonify({"error": "ফাইল দিন"}), 400
        f = request.files['file']
        ext = f.filename.rsplit('.', 1)[-1].lower()
        fp = "/tmp/" + str(uuid.uuid4()) + "." + ext
        f.save(fp)
        m = []
        if ext == 'pdf':
            import pdfplumber
            txt = ""
            with pdfplumber.open(fp) as pdf:
                for pg in pdf.pages:
                    x = pg.extract_text()
                    if x:
                        txt += x + "\n"
            m = pmcqs(txt) if txt.strip() else pimg(fp)
        elif ext in ['png', 'jpg', 'jpeg', 'webp', 'bmp']:
            m = pimg(fp)
        try:
            os.remove(fp)
        except Exception:
            pass
        if not m:
            return jsonify({"error": "MCQ পাওয়া যায়নি"}), 400
        sid = str(uuid.uuid4())[:8]
        sessions[sid] = {'mcqs': m, 'header': {}, 't': time.time()}
        return jsonify({"session_id": sid, "mcqs": m, "count": len(m)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/explain', methods=['POST'])
def api_explain():
    try:
        d = request.get_json()
        sid = d.get('session_id', '')
        subj = d.get('subject', '')
        hdr = d.get('header_info', {})
        if sid not in sessions:
            return jsonify({"error": "সেশন শেষ।"}), 404
        mcqs = sessions[sid]['mcqs']
        sessions[sid]['header'] = hdr
        exp = []
        for i, mcq in enumerate(mcqs):
            r = gexp(mcq, subj)
            exp.append({**mcq, "explanation": r.get("explanation", ""), "sources": r.get("sources", []), "key_point": r.get("key_point", "")})
            if i < len(mcqs) - 1:
                time.sleep(2)
        sessions[sid]['mcqs'] = exp
        return jsonify({"session_id": sid, "mcqs": exp, "count": len(exp)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/export/html', methods=['POST'])
def exp_html():
    d = request.get_json()
    sid = d.get('session_id', '')
    if sid not in sessions:
        return jsonify({"error": "সেশন নেই"}), 404
    s = sessions[sid]
    h = bhtml(s['mcqs'], s.get('header', {}))
    b = io.BytesIO(h.encode('utf-8'))
    b.seek(0)
    return send_file(b, mimetype='text/html', as_attachment=True, download_name="mcq.html")

@app.route('/api/export/word', methods=['POST'])
def exp_word():
    d = request.get_json()
    sid = d.get('session_id', '')
    if sid not in sessions:
        return jsonify({"error": "সেশন নেই"}), 404
    s = sessions[sid]
    b = bword(s['mcqs'], s.get('header', {}))
    return send_file(b, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name="mcq.docx")

@app.route('/api/export/pdf', methods=['POST'])
def exp_pdf():
    d = request.get_json()
    sid = d.get('session_id', '')
    if sid not in sessions:
        return jsonify({"error": "সেশন নেই"}), 404
    s = sessions[sid]
    h = bhtml(s['mcqs'], s.get('header', {}), True)
    b = io.BytesIO(h.encode('utf-8'))
    b.seek(0)
    return send_file(b, mimetype='text/html', as_attachment=True, download_name="mcq_print.html")

def bhtml(mcqs, header, fp=False):
    hh = ""
    if header and any(v for v in header.values() if v):
        hh = "<div class='eh'><h1>" + (header.get('institution', '') or '') + "</h1><h2>" + (header.get('exam_name', '') or '') + "</h2><table class='it'><tr><td><b>সাল:</b> " + header.get('year', '-') + "</td><td><b>পূর্ণমান:</b> " + header.get('total_marks', '-') + "</td><td><b>সময়:</b> " + header.get('duration', '-') + "</td></tr><tr><td><b>নেগেটিভ:</b> " + header.get('negative_marking', '-') + "</td><td colspan='2'><b>নির্দেশনা:</b> " + header.get('instructions', '-') + "</td></tr></table></div>"
    qs = ""
    for i, m in enumerate(mcqs, 1):
        srcs = ', '.join(m.get('sources', [])) or 'N/A'
        ca = (m.get('correct_answer', '') or '').lower()
        om = {'ক': 'a', 'খ': 'b', 'গ': 'c', 'ঘ': 'd', 'a': 'a', 'b': 'b', 'c': 'c', 'd': 'd'}
        cv = om.get(ca, '')
        ac = ' ok' if cv == 'a' else ''
        bc = ' ok' if cv == 'b' else ''
        cc = ' ok' if cv == 'c' else ''
        dc = ' ok' if cv == 'd' else ''
        qs += "<div class='qc'><div class='qn'>প্রশ্ন " + str(i) + "</div><div class='qt'>" + m.get('question', '') + "</div><div class='op" + ac + "'>(ক) " + m.get('option_a', '') + "</div><div class='op" + bc + "'>(খ) " + m.get('option_b', '') + "</div><div class='op" + cc + "'>(গ) " + m.get('option_c', '') + "</div><div class='op" + dc + "'>(ঘ) " + m.get('option_d', '') + "</div><div class='an'>✅ (" + m.get('correct_answer', '') + ") " + m.get('correct_option_text', '') + "</div><div class='ex'><b>📖 ব্যাখ্যা:</b><p>" + m.get('explanation', '') + "</p><div class='sr'><b>📚 তথ্যসূত্র:</b> [" + srcs + "]</div></div></div>"
    pj = '<script>window.onload=function(){window.print()}</script>' if fp else ''
    cs = "*{margin:0;padding:0;box-sizing:border-box}body{font-family:'Noto Sans Bengali',sans-serif;font-size:14px;line-height:1.8;padding:25px;background:#fff}.eh{text-align:center;border:2px solid #1a237e;padding:22px;margin-bottom:28px;border-radius:10px;background:#f8f9ff}.eh h1{font-size:21px;color:#1a237e}.eh h2{font-size:16px;color:#333;margin:5px 0 10px}.it{width:100%;border-collapse:collapse}.it td{padding:6px 12px;border:1px solid #c5cae9;font-size:13px}.qc{border:1px solid #e0e0e0;border-radius:10px;padding:20px;margin-bottom:20px;page-break-inside:avoid;background:#fafbff}.qn{font-weight:700;color:#1565c0;font-size:15px;margin-bottom:8px}.qt{font-size:15px;font-weight:600;margin-bottom:12px}.op{padding:6px 16px;margin:4px 0;border-radius:6px;font-size:14px}.ok{background:#e8f5e9;border-left:4px solid #2e7d32;font-weight:600;color:#1b5e20}.an{background:#e3f2fd;padding:11px 16px;border-radius:7px;margin:14px 0;font-weight:600;color:#0d47a1}.ex{background:#fff;border:1px solid #e0e0e0;padding:16px;border-radius:8px;font-size:13.5px;line-height:1.9}.ex p{margin:8px 0}.sr{margin-top:12px;padding-top:10px;border-top:1px dashed #ccc;font-size:12.5px;color:#555}.ft{text-align:center;margin-top:30px;font-size:12px;color:#999}"
    return "<!DOCTYPE html><html lang='bn'><head><meta charset='UTF-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>MCQ</title><link href='https://fonts.googleapis.com/css2?family=Noto+Sans+Bengali:wght@400;600;700&display=swap' rel='stylesheet'><style>" + cs + "</style>" + pj + "</head><body>" + hh + qs + "<div class='ft'>MCQ সলভার</div></body></html>"

def bword(mcqs, header):
    doc = Document()
    doc.styles['Normal'].font.name = 'Noto Sans Bengali'
    doc.styles['Normal'].font.size = Pt(11)
    if header and any(v for v in header.values() if v):
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(header.get('institution', '') or '')
        r1.bold = True
        r1.font.size = Pt(16)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(header.get('exam_name', '') or '')
        r2.bold = True
        r2.font.size = Pt(13)
    for i, m in enumerate(mcqs, 1):
        pq = doc.add_paragraph()
        rq = pq.add_run("প্রশ্ন " + str(i) + ". " + m.get('question', ''))
        rq.bold = True
        rq.font.size = Pt(12)
        for lb, ky in [('ক', 'option_a'), ('খ', 'option_b'), ('গ', 'option_c'), ('ঘ', 'option_d')]:
            doc.add_paragraph("  (" + lb + ") " + m.get(ky, ''))
        pa = doc.add_paragraph()
        ra = pa.add_run("সঠিক: (" + m.get('correct_answer', '') + ") " + m.get('correct_option_text', ''))
        ra.bold = True
        ra.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        if m.get('explanation'):
            pe = doc.add_paragraph()
            pe.add_run("ব্যাখ্যা: ").bold = True
            pe.add_run(m['explanation'])
        if m.get('sources'):
            ps = doc.add_paragraph()
            rs = ps.add_run("[" + ', '.join(m['sources']) + "]")
            rs.font.size = Pt(9)
            rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph('─' * 40)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
